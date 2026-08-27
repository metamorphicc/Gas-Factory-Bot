import asyncio
import logging
import os
import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from aiogram import Bot, Dispatcher, F, Router
from aiogram.filters import CommandStart
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    CallbackQuery,
    FSInputFile,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    KeyboardButton,
    Message,
    ReplyKeyboardMarkup,
)
from docx import Document
from dotenv import load_dotenv


BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
CLIENTS_DIR = DATA_DIR / "clients"
TEMPLATES_DIR = BASE_DIR / "templates"
OWNER_FILE = DATA_DIR / "owner.txt"

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".jpg", ".jpeg", ".png"}
MAX_FILE_SIZE = 20 * 1024 * 1024
PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*}}")

FIELD_LABELS = {
    "full_name": "ФИО",
    "passport": "Паспорт",
    "inn": "ИНН",
    "date": "Дата",
}

MAIN_BUTTONS = {
    "Добавить клиента",
    "Добавить файл",
    "Получить файл",
    "Заполнить документ",
    "Мои клиенты",
}
NAV_BUTTONS = {"Меню", "В меню"}


class BotStates(StatesGroup):
    waiting_client_name = State()
    add_file_choose_client = State()
    add_file_waiting_file = State()
    get_file_choose_client = State()
    get_file_choose_file = State()
    fill_template_choose_template = State()
    fill_template_choose_client = State()
    fill_template_ask_field = State()
    fill_template_confirm = State()


router = Router()


def ensure_storage() -> None:
    CLIENTS_DIR.mkdir(parents=True, exist_ok=True)
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


def main_menu_keyboard() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="Добавить клиента"), KeyboardButton(text="Добавить файл")],
            [KeyboardButton(text="Получить файл"), KeyboardButton(text="Заполнить документ")],
            [KeyboardButton(text="Мои клиенты")],
        ],
        resize_keyboard=True,
        input_field_placeholder="Выберите действие",
    )


def menu_inline_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[[InlineKeyboardButton(text="В меню", callback_data="menu")]]
    )


def confirmation_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="Все верно", callback_data="fill:ok")],
            [InlineKeyboardButton(text="Отменить", callback_data="menu")],
        ]
    )


def rows_keyboard(prefix: str, items: list[str]) -> InlineKeyboardMarkup:
    rows = [
        [InlineKeyboardButton(text=item, callback_data=f"{prefix}:{index}")]
        for index, item in enumerate(items)
    ]
    rows.append([InlineKeyboardButton(text="В меню", callback_data="menu")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def sanitize_name(raw_name: str) -> str:
    name = raw_name.strip()
    name = re.sub(r"[^\w\s-]", "_", name, flags=re.UNICODE)
    name = re.sub(r"\s+", "_", name)
    name = re.sub(r"_+", "_", name)
    name = name.strip("._-")[:60]
    return name or "client"


def safe_child_path(root: Path, name: str) -> Path:
    candidate = (root / name).resolve()
    root_resolved = root.resolve()
    if root_resolved != candidate and root_resolved not in candidate.parents:
        raise ValueError("Path escapes storage root")
    return candidate


def client_names() -> list[str]:
    ensure_storage()
    return sorted(path.name for path in CLIENTS_DIR.iterdir() if path.is_dir())


def template_names() -> list[str]:
    ensure_storage()
    return sorted(
        path.name
        for path in TEMPLATES_DIR.iterdir()
        if path.is_file() and path.suffix.lower() == ".docx" and not path.name.startswith("~$")
    )


def file_names(client_name: str) -> list[str]:
    folder = safe_child_path(CLIENTS_DIR, client_name)
    return sorted(path.name for path in folder.iterdir() if path.is_file())


def get_template_fields(template_path: Path) -> list[str]:
    found: set[str] = set()
    with zipfile.ZipFile(template_path) as docx_zip:
        for member in docx_zip.namelist():
            if member.startswith("word/") and member.endswith(".xml"):
                text = docx_zip.read(member).decode("utf-8", errors="ignore")
                found.update(PLACEHOLDER_RE.findall(text))

    ordered = [field for field in FIELD_LABELS if field in found]
    ordered.extend(sorted(found - set(ordered)))
    return ordered


def replace_in_paragraph(paragraph: Any, values: dict[str, str]) -> None:
    original = paragraph.text
    updated = original
    for field, value in values.items():
        updated = re.sub(r"{{\s*" + re.escape(field) + r"\s*}}", value, updated)
    if updated != original:
        paragraph.text = updated


def fill_docx(template_path: Path, output_path: Path, values: dict[str, str]) -> None:
    document = Document(template_path)
    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, values)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, values)
    document.save(output_path)


def unique_path(folder: Path, filename: str) -> Path:
    sanitized = sanitize_filename(filename)
    target = safe_child_path(folder, sanitized)
    if not target.exists():
        return target

    stem = target.stem
    suffix = target.suffix
    for number in range(1, 1000):
        candidate = safe_child_path(folder, f"{stem}_{number}{suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError("Too many files with the same name")


def sanitize_filename(raw_filename: str) -> str:
    path = Path(raw_filename)
    stem = sanitize_name(path.stem)
    suffix = path.suffix.lower()
    return f"{stem}{suffix}" if suffix else stem


def mask_value(field: str, value: str) -> str:
    if field not in {"passport", "inn"}:
        return value
    digits = re.sub(r"\D", "", value)
    tail = digits[-4:] if len(digits) >= 4 else value[-2:]
    return f"***{tail}"


async def is_owner(message_or_callback: Message | CallbackQuery) -> bool:
    user = message_or_callback.from_user
    if user is None:
        return False

    user_id = str(user.id)
    ensure_storage()
    if not OWNER_FILE.exists():
        OWNER_FILE.write_text(user_id, encoding="utf-8")
        return True

    owner_id = OWNER_FILE.read_text(encoding="utf-8").strip()
    return owner_id == user_id


async def reject_if_not_owner(message_or_callback: Message | CallbackQuery) -> bool:
    if await is_owner(message_or_callback):
        return False

    text = "Этот MVP доступен только пользователю, который первым запустил демо-сессию."
    if isinstance(message_or_callback, CallbackQuery):
        await message_or_callback.answer("Доступ закрыт", show_alert=True)
        if message_or_callback.message:
            await message_or_callback.message.answer(text)
    else:
        await message_or_callback.answer(text)
    return True


async def show_main_menu(message: Message, state: FSMContext) -> None:
    await state.clear()
    await message.answer("Главное меню", reply_markup=main_menu_keyboard())


async def show_client_list(message: Message) -> None:
    clients = client_names()
    if not clients:
        await message.answer("Клиентов пока нет.", reply_markup=main_menu_keyboard())
        return

    text = "Мои клиенты:\n" + "\n".join(f"- {client}" for client in clients)
    await message.answer(text, reply_markup=main_menu_keyboard())


async def show_clients_for_flow(message: Message, state: FSMContext, flow: str) -> None:
    clients = client_names()
    if not clients:
        await message.answer(
            "Пока нет клиентов. Сначала добавьте клиента.",
            reply_markup=menu_inline_keyboard(),
        )
        return

    await state.update_data(clients=clients)
    if flow == "add":
        await state.set_state(BotStates.add_file_choose_client)
        title = "Выберите клиента, куда сохранить файл:"
        prefix = "add_client"
    elif flow == "get":
        await state.set_state(BotStates.get_file_choose_client)
        title = "Выберите клиента:"
        prefix = "get_client"
    else:
        await state.set_state(BotStates.fill_template_choose_client)
        title = "Выберите клиента для готового документа:"
        prefix = "fill_client"

    await message.answer(title, reply_markup=rows_keyboard(prefix, clients))


@router.message(CommandStart())
async def start(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return
    await show_main_menu(message, state)


@router.callback_query(F.data == "menu")
async def callback_menu(callback: CallbackQuery, state: FSMContext) -> None:
    if await reject_if_not_owner(callback):
        return
    await state.clear()
    await callback.answer()
    if callback.message:
        await callback.message.answer("Главное меню", reply_markup=main_menu_keyboard())


@router.message(F.text.in_(MAIN_BUTTONS | NAV_BUTTONS))
async def main_menu_buttons(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return

    await state.clear()
    if message.text in NAV_BUTTONS:
        await message.answer("Главное меню", reply_markup=main_menu_keyboard())
    elif message.text == "Добавить клиента":
        await state.set_state(BotStates.waiting_client_name)
        await message.answer(
            "Введите имя или короткий ID клиента:",
            reply_markup=menu_inline_keyboard(),
        )
    elif message.text == "Добавить файл":
        await show_clients_for_flow(message, state, "add")
    elif message.text == "Получить файл":
        await show_clients_for_flow(message, state, "get")
    elif message.text == "Заполнить документ":
        await fill_template_start(message, state)
    elif message.text == "Мои клиенты":
        await show_client_list(message)


@router.message(F.text == "Добавить клиента")
async def add_client_start(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return
    await state.set_state(BotStates.waiting_client_name)
    await message.answer(
        "Введите имя или короткий ID клиента:",
        reply_markup=menu_inline_keyboard(),
    )


@router.message(BotStates.waiting_client_name)
async def add_client_finish(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return
    if not message.text:
        await message.answer("Пришлите имя текстом.", reply_markup=menu_inline_keyboard())
        return

    client_name = sanitize_name(message.text)
    client_path = safe_child_path(CLIENTS_DIR, client_name)
    if client_path.exists():
        await message.answer(
            f"Клиент уже есть: data/clients/{client_name}",
            reply_markup=main_menu_keyboard(),
        )
    else:
        client_path.mkdir(parents=True, exist_ok=False)
        await message.answer(
            f"Клиент добавлен: data/clients/{client_name}",
            reply_markup=main_menu_keyboard(),
        )
    await state.clear()


@router.message(F.text == "Мои клиенты")
async def my_clients(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return
    await state.clear()
    clients = client_names()
    if not clients:
        await message.answer("Клиентов пока нет.", reply_markup=main_menu_keyboard())
        return

    text = "Мои клиенты:\n" + "\n".join(f"- {client}" for client in clients)
    await message.answer(text, reply_markup=main_menu_keyboard())


@router.message(F.text == "Добавить файл")
async def add_file_start(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return
    await show_clients_for_flow(message, state, "add")


@router.callback_query(BotStates.add_file_choose_client, F.data.startswith("add_client:"))
async def add_file_choose_client(callback: CallbackQuery, state: FSMContext) -> None:
    if await reject_if_not_owner(callback):
        return
    data = await state.get_data()
    clients = data.get("clients", [])
    index = int(callback.data.split(":", 1)[1])
    if index >= len(clients):
        await callback.answer("Клиент не найден", show_alert=True)
        return

    await state.update_data(selected_client=clients[index])
    await state.set_state(BotStates.add_file_waiting_file)
    await callback.answer()
    if callback.message:
        await callback.message.answer(
            "Пришлите файл: pdf, doc, docx, jpg или png. Максимум 20 МБ.",
            reply_markup=menu_inline_keyboard(),
        )


@router.message(BotStates.add_file_waiting_file)
async def add_file_receive(message: Message, state: FSMContext, bot: Bot) -> None:
    if await reject_if_not_owner(message):
        return

    telegram_file_id: str | None = None
    original_name: str | None = None
    file_size = 0

    if message.document:
        telegram_file_id = message.document.file_id
        original_name = message.document.file_name or "file"
        file_size = message.document.file_size or 0
    elif message.photo:
        photo = message.photo[-1]
        telegram_file_id = photo.file_id
        original_name = f"photo_{photo.file_unique_id}.jpg"
        file_size = photo.file_size or 0

    if not telegram_file_id or not original_name:
        await message.answer("Пришлите файл или фото.", reply_markup=menu_inline_keyboard())
        return

    if file_size > MAX_FILE_SIZE:
        await message.answer("Файл слишком большой. Максимум 20 МБ.", reply_markup=menu_inline_keyboard())
        return

    extension = Path(original_name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        await message.answer("Можно загрузить только pdf, doc, docx, jpg или png.", reply_markup=menu_inline_keyboard())
        return

    data = await state.get_data()
    client_name = data["selected_client"]
    client_folder = safe_child_path(CLIENTS_DIR, client_name)
    target_path = unique_path(client_folder, original_name)

    telegram_file = await bot.get_file(telegram_file_id)
    remote_size = telegram_file.file_size or file_size
    if remote_size > MAX_FILE_SIZE:
        await message.answer("Файл слишком большой. Максимум 20 МБ.", reply_markup=menu_inline_keyboard())
        return

    await bot.download_file(telegram_file.file_path, destination=target_path)
    await state.clear()
    relative_path = target_path.relative_to(BASE_DIR).as_posix()
    await message.answer(f"Файл сохранен: {relative_path}", reply_markup=main_menu_keyboard())


@router.message(F.text == "Получить файл")
async def get_file_start(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return
    await show_clients_for_flow(message, state, "get")


@router.callback_query(BotStates.get_file_choose_client, F.data.startswith("get_client:"))
async def get_file_choose_client(callback: CallbackQuery, state: FSMContext) -> None:
    if await reject_if_not_owner(callback):
        return
    data = await state.get_data()
    clients = data.get("clients", [])
    index = int(callback.data.split(":", 1)[1])
    if index >= len(clients):
        await callback.answer("Клиент не найден", show_alert=True)
        return

    client_name = clients[index]
    files = file_names(client_name)
    if not files:
        await callback.answer()
        if callback.message:
            await callback.message.answer("В папке клиента пока нет файлов.", reply_markup=menu_inline_keyboard())
        return

    await state.update_data(selected_client=client_name, files=files)
    await state.set_state(BotStates.get_file_choose_file)
    await callback.answer()
    if callback.message:
        await callback.message.answer("Выберите файл:", reply_markup=rows_keyboard("get_file", files))


@router.callback_query(BotStates.get_file_choose_file, F.data.startswith("get_file:"))
async def get_file_send(callback: CallbackQuery, state: FSMContext) -> None:
    if await reject_if_not_owner(callback):
        return
    data = await state.get_data()
    files = data.get("files", [])
    index = int(callback.data.split(":", 1)[1])
    if index >= len(files):
        await callback.answer("Файл не найден", show_alert=True)
        return

    client_name = data["selected_client"]
    file_path = safe_child_path(safe_child_path(CLIENTS_DIR, client_name), files[index])
    await state.clear()
    await callback.answer()
    if callback.message:
        await callback.message.answer_document(
            FSInputFile(file_path),
            caption=f"Файл: {files[index]}",
            reply_markup=main_menu_keyboard(),
        )


@router.message(F.text == "Заполнить документ")
async def fill_template_start(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return

    templates = template_names()
    if not templates:
        await message.answer("В папке templates нет .docx шаблонов.", reply_markup=menu_inline_keyboard())
        return

    await state.update_data(templates=templates)
    await state.set_state(BotStates.fill_template_choose_template)
    await message.answer("Выберите шаблон:", reply_markup=rows_keyboard("template", templates))


@router.callback_query(BotStates.fill_template_choose_template, F.data.startswith("template:"))
async def fill_template_choose_template(callback: CallbackQuery, state: FSMContext) -> None:
    if await reject_if_not_owner(callback):
        return
    data = await state.get_data()
    templates = data.get("templates", [])
    index = int(callback.data.split(":", 1)[1])
    if index >= len(templates):
        await callback.answer("Шаблон не найден", show_alert=True)
        return

    template_name = templates[index]
    template_path = safe_child_path(TEMPLATES_DIR, template_name)
    fields = get_template_fields(template_path)
    if not fields:
        await callback.answer()
        if callback.message:
            await callback.message.answer("В шаблоне не найдены поля вида {{field}}.", reply_markup=menu_inline_keyboard())
        return

    await state.update_data(selected_template=template_name, fields=fields, values={})
    await callback.answer()
    if callback.message:
        await show_clients_for_flow(callback.message, state, "fill")


@router.callback_query(BotStates.fill_template_choose_client, F.data.startswith("fill_client:"))
async def fill_template_choose_client(callback: CallbackQuery, state: FSMContext) -> None:
    if await reject_if_not_owner(callback):
        return
    data = await state.get_data()
    clients = data.get("clients", [])
    index = int(callback.data.split(":", 1)[1])
    if index >= len(clients):
        await callback.answer("Клиент не найден", show_alert=True)
        return

    fields = data["fields"]
    first_field = fields[0]
    await state.update_data(selected_client=clients[index], current_field_index=0)
    await state.set_state(BotStates.fill_template_ask_field)
    await callback.answer()
    if callback.message:
        await callback.message.answer(
            f"Введите поле: {FIELD_LABELS.get(first_field, first_field)}",
            reply_markup=menu_inline_keyboard(),
        )


@router.message(BotStates.fill_template_ask_field)
async def fill_template_collect_field(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return
    if not message.text:
        await message.answer("Введите значение текстом.", reply_markup=menu_inline_keyboard())
        return

    data = await state.get_data()
    fields = data["fields"]
    index = data["current_field_index"]
    values = data.get("values", {})
    values[fields[index]] = message.text.strip()
    index += 1

    if index < len(fields):
        await state.update_data(values=values, current_field_index=index)
        next_field = fields[index]
        await message.answer(
            f"Введите поле: {FIELD_LABELS.get(next_field, next_field)}",
            reply_markup=menu_inline_keyboard(),
        )
        return

    await state.update_data(values=values)
    await state.set_state(BotStates.fill_template_confirm)
    lines = [
        "Проверьте данные:",
        f"Клиент: {data['selected_client']}",
        f"Шаблон: {data['selected_template']}",
    ]
    for field in fields:
        label = FIELD_LABELS.get(field, field)
        lines.append(f"{label}: {mask_value(field, values[field])}")
    await message.answer("\n".join(lines), reply_markup=confirmation_keyboard())


@router.callback_query(BotStates.fill_template_confirm, F.data == "fill:ok")
async def fill_template_finish(callback: CallbackQuery, state: FSMContext) -> None:
    if await reject_if_not_owner(callback):
        return

    data = await state.get_data()
    client_name = data["selected_client"]
    template_name = data["selected_template"]
    values = data["values"]

    client_folder = safe_child_path(CLIENTS_DIR, client_name)
    template_path = safe_child_path(TEMPLATES_DIR, template_name)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"{Path(template_name).stem}_{timestamp}.docx"
    output_path = unique_path(client_folder, output_name)
    fill_docx(template_path, output_path, values)

    await state.clear()
    await callback.answer()
    if callback.message:
        relative_path = output_path.relative_to(BASE_DIR).as_posix()
        await callback.message.answer(f"Документ готов: {relative_path}", reply_markup=main_menu_keyboard())
        await callback.message.answer_document(FSInputFile(output_path), caption="Готовый документ")


@router.message(F.text.in_({"Меню", "В меню"}))
async def menu_text(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return
    await show_main_menu(message, state)


@router.message()
async def fallback(message: Message, state: FSMContext) -> None:
    if await reject_if_not_owner(message):
        return
    await message.answer("Выберите действие кнопкой в меню.", reply_markup=main_menu_keyboard())


async def main() -> None:
    load_dotenv()
    token = os.getenv("BOT_TOKEN")
    if not token:
        raise RuntimeError("BOT_TOKEN is missing. Put it into .env")

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s %(levelname)s %(name)s: %(message)s",
    )
    logging.getLogger("aiogram.event").setLevel(logging.WARNING)
    ensure_storage()

    bot = Bot(token=token)
    dispatcher = Dispatcher(storage=MemoryStorage())
    dispatcher.include_router(router)
    await dispatcher.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())
